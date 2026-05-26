#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RAG (Retrieval-Augmented Generation) 知识库核心模块

提供完整的知识库管理、文档解析、向量化存储和检索功能。
"""

import os
import re
import json
import sqlite3
import logging
import hashlib
import datetime
import numpy as np
from typing import List, Dict, Optional, Tuple, Any, Union
from dataclasses import dataclass, asdict
from abc import ABC, abstractmethod
import requests

from api_encryptor import decrypt_api_key, encrypt_api_key
from db_router import DatabaseRouter, DatabasePermission

# 配置日志
logger = logging.getLogger(__name__)

# 数据库文件路径
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
DATA_DIR = os.path.join(BASE_DIR, 'data')
DB_PATH = os.path.join(DATA_DIR, 'rag.db')

# 确保数据目录存在
os.makedirs(DATA_DIR, exist_ok=True)


@dataclass
class KnowledgeBase:
    """知识库数据类"""
    id: Optional[int]
    name: str
    description: str
    embedding_provider: str  # 'deepseek', 'siliconflow', 'openai'
    embedding_model: str
    encrypted_api_key: Optional[str]
    chunk_size: int
    chunk_overlap: int
    created_at: Optional[str] = None
    updated_at: Optional[str] = None


@dataclass
class Document:
    """文档数据类"""
    id: Optional[int]
    kb_id: int
    filename: str
    file_path: str
    file_type: str  # 'pdf', 'docx', 'txt', 'md'
    file_size: int
    content_hash: str
    status: str  # 'pending', 'processing', 'completed', 'failed'
    error_message: Optional[str] = None
    created_at: Optional[str] = None
    updated_at: Optional[str] = None


@dataclass
class Chunk:
    """文本块数据类"""
    id: Optional[int]
    doc_id: int
    kb_id: int
    content: str
    chunk_index: int
    token_count: int
    embedding: Optional[List[float]] = None
    created_at: Optional[str] = None


class RAGDatabase:
    """
    RAG数据库管理类
    
    负责SQLite数据库的初始化和连接管理，使用DatabaseRouter进行安全访问控制。
    """
    
    def __init__(self, db_path: str = DB_PATH):
        self.db_path = db_path
        self._local = threading.local()
        self._router = DatabaseRouter(db_path)
        self._register_whitelist_queries()
        self._init_database()
    
    def _register_whitelist_queries(self) -> None:
        """注册所有需要的查询到白名单"""
        # CREATE TABLE 语句
        self._router.register_whitelist_query('''
            CREATE TABLE IF NOT EXISTS knowledge_bases (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE,
                description TEXT,
                embedding_provider TEXT NOT NULL,
                embedding_model TEXT NOT NULL,
                encrypted_api_key TEXT,
                chunk_size INTEGER DEFAULT 500,
                chunk_overlap INTEGER DEFAULT 50,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        self._router.register_whitelist_query('''
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                kb_id INTEGER NOT NULL,
                filename TEXT NOT NULL,
                file_path TEXT NOT NULL,
                file_type TEXT NOT NULL,
                file_size INTEGER NOT NULL,
                content_hash TEXT NOT NULL,
                status TEXT DEFAULT 'pending',
                error_message TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (kb_id) REFERENCES knowledge_bases(id) ON DELETE CASCADE
            )
        ''')
        
        self._router.register_whitelist_query('''
            CREATE TABLE IF NOT EXISTS chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                doc_id INTEGER NOT NULL,
                kb_id INTEGER NOT NULL,
                content TEXT NOT NULL,
                chunk_index INTEGER NOT NULL,
                token_count INTEGER DEFAULT 0,
                embedding TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (doc_id) REFERENCES documents(id) ON DELETE CASCADE,
                FOREIGN KEY (kb_id) REFERENCES knowledge_bases(id) ON DELETE CASCADE
            )
        ''')
        
        # CREATE INDEX 语句
        self._router.register_whitelist_query('CREATE INDEX IF NOT EXISTS idx_chunks_kb_id ON chunks(kb_id)')
        self._router.register_whitelist_query('CREATE INDEX IF NOT EXISTS idx_chunks_doc_id ON chunks(doc_id)')
        self._router.register_whitelist_query('CREATE INDEX IF NOT EXISTS idx_documents_kb_id ON documents(kb_id)')
        self._router.register_whitelist_query('CREATE INDEX IF NOT EXISTS idx_documents_status ON documents(status)')
        
        # SELECT 查询模板
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+\*\s+FROM\s+knowledge_bases\s*$')
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+\*\s+FROM\s+knowledge_bases\s+WHERE\s+id\s*=\s*\?\s*$')
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+\*\s+FROM\s+knowledge_bases\s+WHERE\s+name\s*=\s*\?\s*$')
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+\*\s+FROM\s+knowledge_bases\s+ORDER\s+BY\s+created_at\s+DESC\s*$')
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+encrypted_api_key\s+FROM\s+knowledge_bases\s+WHERE\s+id\s*=\s*\?\s*$')
        
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+\*\s+FROM\s+documents\s+WHERE\s+kb_id\s*=\s*\?\s+AND\s+content_hash\s*=\s*\?\s*$')
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+\*\s+FROM\s+documents\s+WHERE\s+kb_id\s*=\s*\?\s+ORDER\s+BY\s+created_at\s+DESC\s*$')
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+\*\s+FROM\s+documents\s+WHERE\s+id\s*=\s*\?\s*$')
        
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+id\s*,\s*doc_id\s*,\s*content\s*,\s*embedding\s+FROM\s+chunks\s+WHERE\s+kb_id\s*=\s*\?\s+AND\s+embedding\s+IS\s+NOT\s+NULL\s*$')
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+embedding\s+FROM\s+chunks\s+WHERE\s+id\s*=\s*\?\s*$')
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+COUNT\s*\(\s*\*\s*\)\s+as\s+count\s+FROM\s+chunks\s+WHERE\s+doc_id\s*=\s*\?\s*$')
        self._router.register_whitelist_pattern(r'^\s*SELECT\s+SUM\s*\(\s*token_count\s*\)\s+as\s+total\s+FROM\s+chunks\s+WHERE\s+doc_id\s*=\s*\?\s*$')
        
        # INSERT 语句模板
        self._router.register_whitelist_pattern(r'^\s*INSERT\s+INTO\s+knowledge_bases\s+\([^)]+\)\s+VALUES\s*\([^)]+\)\s*$')
        self._router.register_whitelist_pattern(r'^\s*INSERT\s+INTO\s+documents\s+\([^)]+\)\s+VALUES\s*\([^)]+\)\s*$')
        self._router.register_whitelist_pattern(r'^\s*INSERT\s+INTO\s+chunks\s+\([^)]+\)\s+VALUES\s*\([^)]+\)\s*$')
        
        # UPDATE 语句模板
        self._router.register_whitelist_pattern(r'^\s*UPDATE\s+knowledge_bases\s+SET\s+.+\s+WHERE\s+id\s*=\s*\?\s*$')
        self._router.register_whitelist_pattern(r'^\s*UPDATE\s+documents\s+SET\s+status\s*=\s*\?\s*,\s*updated_at\s*=\s*CURRENT_TIMESTAMP\s+WHERE\s+id\s*=\s*\?\s*$')
        self._router.register_whitelist_pattern(r'^\s*UPDATE\s+documents\s+SET\s+status\s*=\s*\?\s*,\s*error_message\s*=\s*\?\s*,\s*updated_at\s*=\s*CURRENT_TIMESTAMP\s+WHERE\s+id\s*=\s*\?\s*$')
        
        # DELETE 语句模板
        self._router.register_whitelist_pattern(r'^\s*DELETE\s+FROM\s+knowledge_bases\s+WHERE\s+id\s*=\s*\?\s*$')
        self._router.register_whitelist_pattern(r'^\s*DELETE\s+FROM\s+documents\s+WHERE\s+id\s*=\s*\?\s*$')
    
    def _get_connection(self, permission: DatabasePermission = DatabasePermission.READ_WRITE):
        """获取线程本地安全数据库连接
        
        Args:
            permission: 访问权限级别，默认为READ_WRITE
        """
        if not hasattr(self._local, 'connection') or self._local.connection is None:
            self._local.connection = self._router.get_connection(permission)
        return self._local.connection
    
    def _init_database(self) -> None:
        """初始化数据库表结构"""
        try:
            conn = self._get_connection(DatabasePermission.ADMIN)
            
            # 知识库表
            conn.execute('''
                CREATE TABLE IF NOT EXISTS knowledge_bases (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT NOT NULL UNIQUE,
                    description TEXT,
                    embedding_provider TEXT NOT NULL,
                    embedding_model TEXT NOT NULL,
                    encrypted_api_key TEXT,
                    chunk_size INTEGER DEFAULT 500,
                    chunk_overlap INTEGER DEFAULT 50,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            # 文档表
            conn.execute('''
                CREATE TABLE IF NOT EXISTS documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    kb_id INTEGER NOT NULL,
                    filename TEXT NOT NULL,
                    file_path TEXT NOT NULL,
                    file_type TEXT NOT NULL,
                    file_size INTEGER NOT NULL,
                    content_hash TEXT NOT NULL,
                    status TEXT DEFAULT 'pending',
                    error_message TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (kb_id) REFERENCES knowledge_bases(id) ON DELETE CASCADE
                )
            ''')
            
            # 文本块表 - 存储向量化后的文本块
            conn.execute('''
                CREATE TABLE IF NOT EXISTS chunks (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    doc_id INTEGER NOT NULL,
                    kb_id INTEGER NOT NULL,
                    content TEXT NOT NULL,
                    chunk_index INTEGER NOT NULL,
                    token_count INTEGER DEFAULT 0,
                    embedding TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (doc_id) REFERENCES documents(id) ON DELETE CASCADE,
                    FOREIGN KEY (kb_id) REFERENCES knowledge_bases(id) ON DELETE CASCADE
                )
            ''')
            
            # 创建索引优化查询性能
            conn.execute('CREATE INDEX IF NOT EXISTS idx_chunks_kb_id ON chunks(kb_id)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_chunks_doc_id ON chunks(doc_id)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_documents_kb_id ON documents(kb_id)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_documents_status ON documents(status)')
            
            conn.commit()
            logger.info("RAG数据库初始化完成")
            
        except Exception as e:
            logger.error(f"数据库初始化失败: {e}")
            raise
    
    def execute(self, query: str, params: Tuple = (), permission: DatabasePermission = DatabasePermission.READ_WRITE) -> sqlite3.Cursor:
        """执行SQL查询
        
        Args:
            query: SQL查询字符串
            params: 查询参数
            permission: 访问权限级别，默认为READ_WRITE
        """
        conn = self._get_connection(permission)
        cursor = conn.execute(query, params)
        conn.commit()
        return cursor
    
    def fetchone(self, query: str, params: Tuple = (), permission: DatabasePermission = DatabasePermission.READ_WRITE) -> Optional[sqlite3.Row]:
        """获取单条记录
        
        Args:
            query: SQL查询字符串
            params: 查询参数
            permission: 访问权限级别，默认为READ_WRITE
        """
        cursor = self.execute(query, params, permission)
        return cursor.fetchone()
    
    def fetchall(self, query: str, params: Tuple = (), permission: DatabasePermission = DatabasePermission.READ_WRITE) -> List[sqlite3.Row]:
        """获取多条记录
        
        Args:
            query: SQL查询字符串
            params: 查询参数
            permission: 访问权限级别，默认为READ_WRITE
        """
        cursor = self.execute(query, params, permission)
        return cursor.fetchall()
    
    def close(self) -> None:
        """关闭数据库连接"""
        if hasattr(self._local, 'connection') and self._local.connection:
            self._local.connection.close()
            self._local.connection = None
        self._router.close_all()


import threading


class KnowledgeBaseManager:
    """
    知识库管理类
    
    提供知识库的CRUD操作。
    """
    
    def __init__(self, db: RAGDatabase):
        self.db = db
    
    def create_knowledge_base(
        self,
        name: str,
        description: str,
        embedding_provider: str,
        embedding_model: str,
        api_key: Optional[str] = None,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ) -> Optional[int]:
        """
        创建新知识库
        
        Args:
            name: 知识库名称
            description: 知识库描述
            embedding_provider: 嵌入服务提供商
            embedding_model: 嵌入模型名称
            api_key: API密钥（会被加密存储）
            chunk_size: 分块大小
            chunk_overlap: 分块重叠大小
        
        Returns:
            新创建的知识库ID，失败返回None
        """
        try:
            encrypted_key = encrypt_api_key(api_key) if api_key else None
            
            cursor = self.db.execute('''
                INSERT INTO knowledge_bases 
                (name, description, embedding_provider, embedding_model, 
                 encrypted_api_key, chunk_size, chunk_overlap)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (name, description, embedding_provider, embedding_model, 
                  encrypted_key, chunk_size, chunk_overlap))
            
            kb_id = cursor.lastrowid
            logger.info(f"创建知识库成功: {name} (ID: {kb_id})")
            return kb_id
            
        except sqlite3.IntegrityError:
            logger.error(f"知识库名称已存在: {name}")
            return None
        except Exception as e:
            logger.error(f"创建知识库失败: {e}")
            return None
    
    def get_knowledge_base(self, kb_id: int) -> Optional[KnowledgeBase]:
        """根据ID获取知识库"""
        row = self.db.fetchone(
            'SELECT * FROM knowledge_bases WHERE id = ?', (kb_id,)
        )
        if row:
            return KnowledgeBase(**dict(row))
        return None
    
    def get_knowledge_base_by_name(self, name: str) -> Optional[KnowledgeBase]:
        """根据名称获取知识库"""
        row = self.db.fetchone(
            'SELECT * FROM knowledge_bases WHERE name = ?', (name,)
        )
        if row:
            return KnowledgeBase(**dict(row))
        return None
    
    def list_knowledge_bases(self) -> List[KnowledgeBase]:
        """获取所有知识库列表"""
        rows = self.db.fetchall('SELECT * FROM knowledge_bases ORDER BY created_at DESC')
        return [KnowledgeBase(**dict(row)) for row in rows]
    
    def update_knowledge_base(self, kb_id: int, **kwargs) -> bool:
        """更新知识库信息"""
        try:
            allowed_fields = ['name', 'description', 'embedding_model', 'chunk_size', 'chunk_overlap']
            updates = {k: v for k, v in kwargs.items() if k in allowed_fields}
            
            if not updates:
                return False
            
            updates['updated_at'] = datetime.datetime.now().isoformat()
            
            set_clause = ', '.join([f"{k} = ?" for k in updates.keys()])
            values = list(updates.values()) + [kb_id]
            
            self.db.execute(
                f'UPDATE knowledge_bases SET {set_clause} WHERE id = ?',
                tuple(values)
            )
            logger.info(f"更新知识库成功: ID={kb_id}")
            return True
            
        except Exception as e:
            logger.error(f"更新知识库失败: {e}")
            return False
    
    def delete_knowledge_base(self, kb_id: int) -> bool:
        """删除知识库及其所有关联数据"""
        try:
            self.db.execute('DELETE FROM knowledge_bases WHERE id = ?', (kb_id,))
            logger.info(f"删除知识库成功: ID={kb_id}")
            return True
        except Exception as e:
            logger.error(f"删除知识库失败: {e}")
            return False
    
    def get_api_key(self, kb_id: int) -> Optional[str]:
        """获取知识库的解密API密钥"""
        row = self.db.fetchone(
            'SELECT encrypted_api_key FROM knowledge_bases WHERE id = ?', (kb_id,)
        )
        if row and row['encrypted_api_key']:
            return decrypt_api_key(row['encrypted_api_key'])
        return None


class DocumentParser:
    """
    文档解析器
    
    支持PDF、Word、TXT、Markdown等格式的文档解析。
    """
    
    SUPPORTED_TYPES = {'.pdf', '.docx', '.txt', '.md', '.markdown'}
    
    @classmethod
    def parse(cls, file_path: str) -> Tuple[str, str]:
        """
        解析文档内容
        
        Args:
            file_path: 文档文件路径
        
        Returns:
            (content, file_type) 元组
        
        Raises:
            ValueError: 不支持的文件类型
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in cls.SUPPORTED_TYPES:
            raise ValueError(f"不支持的文件类型: {ext}")
        
        try:
            if ext == '.pdf':
                return cls._parse_pdf(file_path), 'pdf'
            elif ext == '.docx':
                return cls._parse_docx(file_path), 'docx'
            elif ext in {'.txt', '.md', '.markdown'}:
                return cls._parse_text(file_path), 'txt'
            else:
                raise ValueError(f"未实现的文件类型解析: {ext}")
                
        except Exception as e:
            logger.error(f"解析文档失败 {file_path}: {e}")
            raise
    
    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """解析PDF文档"""
        try:
            import PyPDF2
            text_parts = []
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            return '\n\n'.join(text_parts)
            
        except ImportError:
            logger.error("PyPDF2未安装，无法解析PDF")
            raise
        except Exception as e:
            logger.error(f"PDF解析错误: {e}")
            raise
    
    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """解析Word文档"""
        try:
            import docx
            doc = docx.Document(file_path)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
            
        except ImportError:
            logger.error("python-docx未安装，无法解析Word文档")
            raise
        except Exception as e:
            logger.error(f"Word解析错误: {e}")
            raise
    
    @staticmethod
    def _parse_text(file_path: str) -> str:
        """解析纯文本文档"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"无法识别文件编码: {file_path}")
    
    @staticmethod
    def compute_hash(file_path: str) -> str:
        """计算文件内容的MD5哈希"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()


class TextSplitter:
    """
    递归字符分块器
    
    实现递归字符文本分割，优先按段落、句子分割，保持语义完整性。
    """
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        separators: Optional[List[str]] = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", "。", "，", " ", ""]
    
    def split_text(self, text: str) -> List[str]:
        """
        将文本分割成块
        
        Args:
            text: 输入文本
        
        Returns:
            文本块列表
        """
        if not text:
            return []
        
        chunks = self._recursive_split(text, 0)
        return chunks
    
    def _recursive_split(self, text: str, separator_index: int) -> List[str]:
        """递归分割文本"""
        separator = self.separators[separator_index]
        
        # 如果没有分隔符或已是最后一个，直接按字符分割
        if separator == "" or separator_index >= len(self.separators) - 1:
            return self._split_by_chars(text)
        
        # 按当前分隔符分割
        parts = text.split(separator)
        
        chunks = []
        current_chunk = ""
        
        for part in parts:
            part = part.strip()
            if not part:
                continue
            
            # 如果当前部分加上分隔符后超过块大小
            if len(current_chunk) + len(part) + len(separator) > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # 如果单个部分超过块大小，递归使用更细粒度的分隔符
                if len(part) > self.chunk_size:
                    sub_chunks = self._recursive_split(part, separator_index + 1)
                    chunks.extend(sub_chunks)
                    current_chunk = ""
                else:
                    current_chunk = part
            else:
                if current_chunk:
                    current_chunk += separator + part
                else:
                    current_chunk = part
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # 添加重叠
        return self._add_overlap(chunks)
    
    def _split_by_chars(self, text: str) -> List[str]:
        """按字符数分割文本"""
        chunks = []
        for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
            chunk = text[i:i + self.chunk_size]
            if chunk:
                chunks.append(chunk)
        return chunks
    
    def _add_overlap(self, chunks: List[str]) -> List[str]:
        """在块之间添加重叠内容"""
        if len(chunks) <= 1 or self.chunk_overlap <= 0:
            return chunks
        
        overlapped_chunks = [chunks[0]]
        
        for i in range(1, len(chunks)):
            prev_chunk = chunks[i - 1]
            current_chunk = chunks[i]
            
            # 从前一块末尾取重叠内容
            overlap_text = prev_chunk[-self.chunk_overlap:] if len(prev_chunk) > self.chunk_overlap else prev_chunk
            
            # 如果当前块不以重叠内容开头，则添加
            if not current_chunk.startswith(overlap_text):
                current_chunk = overlap_text + current_chunk
            
            overlapped_chunks.append(current_chunk)
        
        return overlapped_chunks
    
    def estimate_tokens(self, text: str) -> int:
        """估算文本的token数量（简化版：中文字符 + 英文单词）"""
        # 中文字符
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        # 英文单词
        english_words = len(re.findall(r'[a-zA-Z]+', text))
        
        return chinese_chars + english_words


class EmbeddingClient:
    """
    嵌入向量客户端
    
    支持DeepSeek、SiliconFlow、OpenAI的Embedding API。
    """
    
    PROVIDER_CONFIGS = {
        'deepseek': {
            'url': 'https://api.deepseek.com/v1/embeddings',
            'default_model': 'text-embedding-ada-002'
        },
        'siliconflow': {
            'url': 'https://api.siliconflow.cn/v1/embeddings',
            'default_model': 'BAAI/bge-large-zh-v1.5'
        },
        'openai': {
            'url': 'https://api.openai.com/v1/embeddings',
            'default_model': 'text-embedding-3-small'
        }
    }
    
    def __init__(self, provider: str, api_key: str, model: Optional[str] = None):
        """
        初始化嵌入客户端
        
        Args:
            provider: 提供商名称 ('deepseek', 'siliconflow', 'openai')
            api_key: API密钥
            model: 模型名称（可选，使用默认值）
        """
        if provider not in self.PROVIDER_CONFIGS:
            raise ValueError(f"不支持的嵌入提供商: {provider}")
        
        self.provider = provider
        self.api_key = api_key
        self.config = self.PROVIDER_CONFIGS[provider]
        self.model = model or self.config['default_model']
    
    def embed_texts(self, texts: List[str]) -> List[List[float]]:
        """
        批量获取文本的嵌入向量
        
        Args:
            texts: 文本列表
        
        Returns:
            嵌入向量列表
        """
        if not texts:
            return []
        
        try:
            headers = {
                'Authorization': f'Bearer {self.api_key}',
                'Content-Type': 'application/json'
            }
            
            data = {
                'model': self.model,
                'input': texts,
                'encoding_format': 'float'
            }
            
            response = requests.post(
                self.config['url'],
                headers=headers,
                json=data,
                timeout=60
            )
            response.raise_for_status()
            
            result = response.json()
            embeddings = [item['embedding'] for item in result['data']]
            
            logger.debug(f"成功获取 {len(texts)} 个文本的嵌入向量")
            return embeddings
            
        except requests.exceptions.RequestException as e:
            logger.error(f"嵌入API请求失败: {e}")
            raise
        except (KeyError, IndexError) as e:
            logger.error(f"解析嵌入API响应失败: {e}")
            raise
        except Exception as e:
            logger.error(f"获取嵌入向量失败: {e}")
            raise
    
    def embed_query(self, text: str) -> List[float]:
        """获取单个查询文本的嵌入向量"""
        embeddings = self.embed_texts([text])
        return embeddings[0] if embeddings else []


class VectorRetriever:
    """
    向量检索器
    
    使用numpy计算余弦相似度进行向量检索。
    """
    
    def __init__(self, db: RAGDatabase):
        self.db = db
    
    def retrieve(
        self,
        kb_id: int,
        query_embedding: List[float],
        top_k: int = 5,
        min_score: float = 0.0
    ) -> List[Dict[str, Any]]:
        """
        检索最相似的文本块
        
        Args:
            kb_id: 知识库ID
            query_embedding: 查询向量
            top_k: 返回结果数量
            min_score: 最小相似度阈值
        
        Returns:
            相似文本块列表，包含content, score, doc_id等信息
        """
        try:
            # 获取该知识库的所有文本块
            rows = self.db.fetchall(
                'SELECT id, doc_id, content, embedding FROM chunks WHERE kb_id = ? AND embedding IS NOT NULL',
                (kb_id,)
            )
            
            if not rows:
                logger.warning(f"知识库 {kb_id} 中没有可用的文本块")
                return []
            
            query_vec = np.array(query_embedding, dtype=np.float32)
            
            results = []
            for row in rows:
                if not row['embedding']:
                    continue
                
                try:
                    chunk_embedding = json.loads(row['embedding'])
                    chunk_vec = np.array(chunk_embedding, dtype=np.float32)
                    
                    # 计算余弦相似度
                    similarity = self._cosine_similarity(query_vec, chunk_vec)
                    
                    if similarity >= min_score:
                        results.append({
                            'chunk_id': row['id'],
                            'doc_id': row['doc_id'],
                            'content': row['content'],
                            'score': float(similarity)
                        })
                        
                except (json.JSONDecodeError, ValueError) as e:
                    logger.warning(f"解析嵌入向量失败: {e}")
                    continue
            
            # 按相似度排序并返回top_k
            results.sort(key=lambda x: x['score'], reverse=True)
            return results[:top_k]
            
        except Exception as e:
            logger.error(f"向量检索失败: {e}")
            return []
    
    @staticmethod
    def _cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
        """计算两个向量的余弦相似度"""
        norm_a = np.linalg.norm(a)
        norm_b = np.linalg.norm(b)
        
        if norm_a == 0 or norm_b == 0:
            return 0.0
        
        return float(np.dot(a, b) / (norm_a * norm_b))
    
    def retrieve_with_mmr(
        self,
        kb_id: int,
        query_embedding: List[float],
        top_k: int = 5,
        diversity_lambda: float = 0.5
    ) -> List[Dict[str, Any]]:
        """
        使用MMR(最大边际相关性)算法检索，增加结果多样性
        
        Args:
            kb_id: 知识库ID
            query_embedding: 查询向量
            top_k: 返回结果数量
            diversity_lambda: 多样性权重 (0-1)，越大多样性越高
        
        Returns:
            多样化文本块列表
        """
        # 先获取更多候选结果
        candidates = self.retrieve(kb_id, query_embedding, top_k=top_k * 3, min_score=0.0)
        
        if not candidates:
            return []
        
        query_vec = np.array(query_embedding, dtype=np.float32)
        selected = []
        remaining = candidates.copy()
        
        while len(selected) < top_k and remaining:
            if not selected:
                # 第一个选择最相关的
                best = max(remaining, key=lambda x: x['score'])
            else:
                # MMR评分: λ * Sim(q, d) - (1-λ) * max(Sim(d, d_selected))
                best_mmr_score = -float('inf')
                best = None
                
                for candidate in remaining:
                    relevance = candidate['score']
                    
                    # 计算与已选结果的最大相似度
                    candidate_vec = np.array(json.loads(
                        self.db.fetchone(
                            'SELECT embedding FROM chunks WHERE id = ?',
                            (candidate['chunk_id'],)
                        )['embedding']
                    ), dtype=np.float32)
                    
                    max_sim_to_selected = max(
                        self._cosine_similarity(
                            candidate_vec,
                            np.array(json.loads(
                                self.db.fetchone(
                                    'SELECT embedding FROM chunks WHERE id = ?',
                                    (s['chunk_id'],)
                                )['embedding']
                            ), dtype=np.float32)
                        )
                        for s in selected
                    )
                    
                    mmr_score = diversity_lambda * relevance - (1 - diversity_lambda) * max_sim_to_selected
                    
                    if mmr_score > best_mmr_score:
                        best_mmr_score = mmr_score
                        best = candidate
            
            selected.append(best)
            remaining.remove(best)
        
        return selected


class RAGChat:
    """
    RAG对话构建器
    
    整合检索和LLM调用，实现RAG对话功能。
    """
    
    DEFAULT_SYSTEM_PROMPT = """你是一个基于知识库的智能助手。请根据提供的参考文档回答用户问题。

回答要求：
1. 优先使用参考文档中的信息回答问题
2. 如果参考文档中没有相关信息，请明确说明
3. 保持回答简洁、准确、专业
4. 引用参考文档时，可以标注来源

参考文档：
{context}
"""
    
    def __init__(
        self,
        db: RAGDatabase,
        kb_manager: KnowledgeBaseManager,
        retriever: VectorRetriever
    ):
        self.db = db
        self.kb_manager = kb_manager
        self.retriever = retriever
    
    def chat(
        self,
        kb_id: int,
        query: str,
        conversation_history: Optional[List[Dict[str, str]]] = None,
        top_k: int = 5,
        stream: bool = False
    ) -> Union[str, Any]:
        """
        执行RAG对话
        
        Args:
            kb_id: 知识库ID
            query: 用户查询
            conversation_history: 对话历史
            top_k: 检索文档数量
            stream: 是否流式返回
        
        Returns:
            生成的回答（字符串或生成器）
        """
        try:
            # 获取知识库配置
            kb = self.kb_manager.get_knowledge_base(kb_id)
            if not kb:
                raise ValueError(f"知识库不存在: {kb_id}")
            
            # 获取API密钥
            api_key = self.kb_manager.get_api_key(kb_id)
            if not api_key:
                raise ValueError(f"知识库未配置API密钥: {kb_id}")
            
            # 创建嵌入客户端获取查询向量
            embedding_client = EmbeddingClient(
                kb.embedding_provider,
                api_key,
                kb.embedding_model
            )
            
            query_embedding = embedding_client.embed_query(query)
            
            # 检索相关文档
            retrieved_docs = self.retriever.retrieve(kb_id, query_embedding, top_k=top_k)
            
            if not retrieved_docs:
                context = "（未找到相关文档）"
            else:
                context = "\n\n".join([
                    f"[文档 {i+1}] {doc['content'][:500]}..."
                    for i, doc in enumerate(retrieved_docs)
                ])
            
            # 构建系统提示词
            system_prompt = self.DEFAULT_SYSTEM_PROMPT.format(context=context)
            
            # 构建消息列表
            messages = [{"role": "system", "content": system_prompt}]
            
            if conversation_history:
                messages.extend(conversation_history)
            
            messages.append({"role": "user", "content": query})
            
            # 调用LLM
            if stream:
                return self._call_llm_stream(kb.embedding_provider, api_key, messages)
            else:
                return self._call_llm(kb.embedding_provider, api_key, messages)
                
        except Exception as e:
            logger.error(f"RAG对话失败: {e}")
            raise
    
    def _call_llm(self, provider: str, api_key: str, messages: List[Dict[str, str]]) -> str:
        """调用LLM获取完整回复"""
        # 复用deepseek_parser的API调用逻辑
        from deepseek_parser import DeepSeekWorker
        
        url = "https://api.deepseek.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": messages,
            "temperature": 0.7,
            "max_tokens": 2000
        }
        
        try:
            response = requests.post(url, headers=headers, json=data, timeout=60)
            response.raise_for_status()
            
            result = response.json()
            content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"LLM调用失败: {e}")
            raise
    
    def _call_llm_stream(self, provider: str, api_key: str, messages: List[Dict[str, str]]):
        """流式调用LLM"""
        url = "https://api.deepseek.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": messages,
            "temperature": 0.7,
            "max_tokens": 2000,
            "stream": True
        }
        
        try:
            response = requests.post(url, headers=headers, json=data, timeout=60, stream=True)
            response.raise_for_status()
            
            for line in response.iter_lines():
                if line:
                    line = line.decode('utf-8')
                    if line.startswith('data: '):
                        data_str = line[6:]
                        if data_str == '[DONE]':
                            break
                        try:
                            data_json = json.loads(data_str)
                            delta = data_json.get('choices', [{}])[0].get('delta', {})
                            content = delta.get('content', '')
                            if content:
                                yield content
                        except json.JSONDecodeError:
                            continue
                            
        except Exception as e:
            logger.error(f"流式LLM调用失败: {e}")
            raise


class DocumentProcessor:
    """
    文档处理器
    
    整合文档解析、分块、向量化存储的完整流程。
    """
    
    def __init__(
        self,
        db: RAGDatabase,
        kb_manager: KnowledgeBaseManager,
        retriever: VectorRetriever
    ):
        self.db = db
        self.kb_manager = kb_manager
        self.retriever = retriever
    
    def add_document(self, kb_id: int, file_path: str) -> Optional[int]:
        """
        添加文档到知识库
        
        Args:
            kb_id: 知识库ID
            file_path: 文档文件路径
        
        Returns:
            文档ID，失败返回None
        """
        try:
            # 检查文件是否已存在
            content_hash = DocumentParser.compute_hash(file_path)
            existing = self.db.fetchone(
                'SELECT id FROM documents WHERE kb_id = ? AND content_hash = ?',
                (kb_id, content_hash)
            )
            
            if existing:
                logger.info(f"文档已存在，跳过: {file_path}")
                return existing['id']
            
            # 获取知识库配置
            kb = self.kb_manager.get_knowledge_base(kb_id)
            if not kb:
                raise ValueError(f"知识库不存在: {kb_id}")
            
            # 创建文档记录
            filename = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            
            cursor = self.db.execute('''
                INSERT INTO documents (kb_id, filename, file_path, file_type, file_size, content_hash, status)
                VALUES (?, ?, ?, ?, ?, ?, 'processing')
            ''', (kb_id, filename, file_path, os.path.splitext(file_path)[1][1:], 
                  file_size, content_hash))
            
            doc_id = cursor.lastrowid
            
            # 解析文档
            content, file_type = DocumentParser.parse(file_path)
            
            # 分块
            splitter = TextSplitter(kb.chunk_size, kb.chunk_overlap)
            chunks = splitter.split_text(content)
            
            # 获取API密钥并创建嵌入客户端
            api_key = self.kb_manager.get_api_key(kb_id)
            if not api_key:
                raise ValueError(f"知识库未配置API密钥: {kb_id}")
            
            embedding_client = EmbeddingClient(
                kb.embedding_provider,
                api_key,
                kb.embedding_model
            )
            
            # 批量处理文本块（每批10个）
            batch_size = 10
            for i in range(0, len(chunks), batch_size):
                batch_chunks = chunks[i:i + batch_size]
                
                # 获取嵌入向量
                embeddings = embedding_client.embed_texts(batch_chunks)
                
                # 存储文本块
                for j, (chunk_text, embedding) in enumerate(zip(batch_chunks, embeddings)):
                    token_count = splitter.estimate_tokens(chunk_text)
                    
                    self.db.execute('''
                        INSERT INTO chunks (doc_id, kb_id, content, chunk_index, token_count, embedding)
                        VALUES (?, ?, ?, ?, ?, ?)
                    ''', (doc_id, kb_id, chunk_text, i + j, token_count, 
                          json.dumps(embedding)))
                
                logger.info(f"已处理 {min(i + batch_size, len(chunks))}/{len(chunks)} 个文本块")
            
            # 更新文档状态
            self.db.execute(
                "UPDATE documents SET status = 'completed', updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                (doc_id,)
            )
            
            logger.info(f"文档处理完成: {filename} (ID: {doc_id}), 共 {len(chunks)} 个文本块")
            return doc_id
            
        except Exception as e:
            logger.error(f"添加文档失败: {e}")
            # 更新文档状态为失败
            if 'doc_id' in locals():
                self.db.execute(
                    "UPDATE documents SET status = 'failed', error_message = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                    (str(e)[:500], doc_id)
                )
            return None
    
    def delete_document(self, doc_id: int) -> bool:
        """删除文档及其所有文本块"""
        try:
            self.db.execute('DELETE FROM documents WHERE id = ?', (doc_id,))
            logger.info(f"删除文档成功: ID={doc_id}")
            return True
        except Exception as e:
            logger.error(f"删除文档失败: {e}")
            return False
    
    def list_documents(self, kb_id: int) -> List[Document]:
        """获取知识库的所有文档"""
        rows = self.db.fetchall(
            'SELECT * FROM documents WHERE kb_id = ? ORDER BY created_at DESC',
            (kb_id,)
        )
        return [Document(**dict(row)) for row in rows]
    
    def get_document_stats(self, doc_id: int) -> Dict[str, Any]:
        """获取文档统计信息"""
        doc = self.db.fetchone('SELECT * FROM documents WHERE id = ?', (doc_id,))
        if not doc:
            return {}
        
        chunks_count = self.db.fetchone(
            'SELECT COUNT(*) as count FROM chunks WHERE doc_id = ?',
            (doc_id,)
        )['count']
        
        total_tokens = self.db.fetchone(
            'SELECT SUM(token_count) as total FROM chunks WHERE doc_id = ?',
            (doc_id,)
        )['total'] or 0
        
        return {
            'document': dict(doc),
            'chunks_count': chunks_count,
            'total_tokens': total_tokens
        }


# 便捷函数：创建完整的RAG系统
def create_rag_system() -> Tuple[RAGDatabase, KnowledgeBaseManager, DocumentProcessor, VectorRetriever, RAGChat]:
    """
    创建完整的RAG系统实例
    
    Returns:
        (db, kb_manager, doc_processor, retriever, rag_chat) 元组
    """
    db = RAGDatabase()
    kb_manager = KnowledgeBaseManager(db)
    retriever = VectorRetriever(db)
    doc_processor = DocumentProcessor(db, kb_manager, retriever)
    rag_chat = RAGChat(db, kb_manager, retriever)
    
    return db, kb_manager, doc_processor, retriever, rag_chat


if __name__ == '__main__':
    # 测试代码
    logging.basicConfig(level=logging.INFO)
    
    # 创建RAG系统
    db, kb_manager, doc_processor, retriever, rag_chat = create_rag_system()
    
    # 测试创建知识库
    kb_id = kb_manager.create_knowledge_base(
        name="测试知识库",
        description="用于测试的知识库",
        embedding_provider="deepseek",
        embedding_model="text-embedding-ada-002",
        api_key="your-api-key-here"
    )
    
    print(f"创建知识库ID: {kb_id}")
    
    # 列出所有知识库
    kbs = kb_manager.list_knowledge_bases()
    print(f"知识库列表: {[kb.name for kb in kbs]}")
